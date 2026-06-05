import json
import os
import re
import unicodedata
import uuid
from typing import Dict, List
import fitz
import pytesseract
from docx import Document as DocxDocument
from PIL import Image, ImageFilter, ImageEnhance
import io

from config import GEMINI_API_KEY, LLM_MODEL, OCR_MIN_CHARS_PER_PAGE
from core.genai_client import configure_genai, make_model


class DocumentProcessor:
    """Process raw files into cleaned text and structured fields."""

    def __init__(self):
        configure_genai()
        self.model = make_model(LLM_MODEL)

    def process(self, file_path: str) -> Dict:
        """Detect file type, extract text, clean it, and extract structured fields."""
        file_name = os.path.basename(file_path)
        extension = os.path.splitext(file_name)[1].lower()
        extraction_method = "text"

        if extension == ".pdf":
            pages = self._extract_pdf(file_path)
            extraction_method = "hybrid"
        elif extension == ".docx":
            pages = self._extract_docx(file_path)
            extraction_method = "docx"
        elif extension == ".txt":
            pages = self._extract_txt(file_path)
            extraction_method = "txt"
        else:
            raise ValueError(f"Unsupported file type: {extension}")

        cleaned_pages = []
        for page in pages:
            cleaned_text = self._clean_text(page.get("text", ""))
            cleaned_pages.append({"page_num": page["page_num"], "text": cleaned_text})

        raw_text = "\n\n".join(page["text"] for page in cleaned_pages).strip()
        structured_fields = self._extract_structured_fields(raw_text)

        return {
            "doc_id": str(uuid.uuid4()),
            "file_name": file_name,
            "raw_text": raw_text,
            "pages": cleaned_pages,
            "structured_fields": structured_fields,
            "extraction_method": extraction_method,
        }

    def _extract_pdf(self, file_path: str) -> List[Dict]:
        """Extract text from PDF using hybrid page-level extraction (text + OCR).
        
        For each page:
        - Try native text extraction first
        - Check text quality
        - If poor quality, OCR only that page
        - If good quality, use extracted text
        """
        pages = []
        try:
            document = fitz.open(file_path)
            for page_num in range(len(document)):
                page = document[page_num]
                page_number = page_num + 1
                
                # Try native text extraction first
                native_text = page.get_text().strip()
                
                # Check if text quality is good
                if not self._is_bad_text(native_text):
                    pages.append({
                        "page_num": page_number,
                        "text": native_text,
                        "method": "text",
                        "char_count": len(native_text)
                    })
                else:
                    # Text quality is poor, use OCR for this page
                    try:
                        # Render page to image at 250 DPI
                        pixmap = page.get_pixmap(matrix=fitz.Matrix(250/72, 250/72))
                        
                        # Convert pixmap to PIL Image
                        img_data = pixmap.tobytes("ppm")
                        image = Image.open(io.BytesIO(img_data))
                        
                        # Preprocess image for better OCR
                        processed_image = self._preprocess_image(image)
                        
                        # Run OCR
                        ocr_text = pytesseract.image_to_string(processed_image, lang="eng").strip()
                        
                        pages.append({
                            "page_num": page_number,
                            "text": ocr_text,
                            "method": "ocr",
                            "char_count": len(ocr_text)
                        })
                    except Exception as ocr_error:
                        print(f"⚠️  OCR failed for page {page_number}: {ocr_error}")
                        # Fallback to empty text if OCR fails
                        pages.append({
                            "page_num": page_number,
                            "text": "",
                            "method": "failed",
                            "char_count": 0
                        })
        except Exception as e:
            print(f"❌ Error extracting PDF: {e}")
            pages = [{"page_num": 1, "text": "", "method": "failed", "char_count": 0}]
        
        return pages

    def _is_bad_text(self, text: str) -> bool:
        """Detect if extracted text is of poor quality (scanned or broken extraction)."""
        # Check if empty
        if not text or len(text) == 0:
            return True
        
        # Check minimum character threshold
        if len(text) < OCR_MIN_CHARS_PER_PAGE:
            return True
        
        # Calculate alphabet character ratio
        alphabet_chars = sum(1 for c in text if c.isalpha())
        total_chars = len(text)
        
        if total_chars == 0:
            return True
        
        alphabet_ratio = alphabet_chars / total_chars
        
        # If less than 50% alphabetic characters, text quality is poor
        if alphabet_ratio < 0.5:
            return True
        
        return False

    def _preprocess_image(self, image: Image.Image) -> Image.Image:
        """Preprocess image for better OCR readability."""
        try:
            # Convert to grayscale
            image = image.convert("L")
            
            # Increase contrast
            enhancer = ImageEnhance.Contrast(image)
            image = enhancer.enhance(2.0)
            
            # Sharpen
            image = image.filter(ImageFilter.SHARPEN)
            
            # Increase brightness
            enhancer = ImageEnhance.Brightness(image)
            image = enhancer.enhance(1.1)
        except Exception as e:
            print(f"⚠️  Image preprocessing error: {e}")
            # Return original image if preprocessing fails
        
        return image

    def _extract_docx(self, file_path: str) -> List[Dict]:
        """Extract text from a DOCX file including paragraphs and tables."""
        pages = []
        try:
            document = DocxDocument(file_path)
            all_text = []
            
            # Extract paragraphs
            for paragraph in document.paragraphs:
                if paragraph.text.strip():
                    all_text.append(paragraph.text.strip())
            
            # Extract tables
            for table in document.tables:
                table_text = []
                for row in table.rows:
                    row_cells = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_cells.append(cell.text.strip())
                    if row_cells:
                        table_text.append(" | ".join(row_cells))
                if table_text:
                    all_text.append("\n".join(table_text))
            
            # Split text into pages (every 40 items)
            if all_text:
                page_texts = ["\n".join(all_text[i : i + 40]) for i in range(0, len(all_text), 40)]
            else:
                page_texts = [""]
            
            for index, text in enumerate(page_texts):
                pages.append({"page_num": index + 1, "text": text})
        except Exception as e:
            print(f"❌ Error extracting DOCX: {e}")
            pages = [{"page_num": 1, "text": ""}]
        
        return pages

    def _extract_txt(self, file_path: str) -> List[Dict]:
        """Read plain text files as a single page."""
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as handle:
                text = handle.read().strip()
        except Exception as e:
            print(f"❌ Error reading text file: {e}")
            text = ""
        return [{"page_num": 1, "text": text}]

    def _clean_text(self, text: str) -> str:
        """Normalize whitespace and fix common OCR artifacts while preserving legal formatting."""
        normalized = unicodedata.normalize("NFKC", text)
        normalized = normalized.replace("\r", "\n")
        # Preserve paragraph separation for legal documents (max 2 newlines)
        normalized = re.sub(r"\n{3,}", "\n\n", normalized)
        # Fix excessive spaces but preserve intentional indentation
        normalized = re.sub(r"[ \t]{2,}", " ", normalized)
        # Fix common OCR mistakes
        normalized = normalized.replace("c1ause", "clause").replace("lndemnification", "indemnification")
        normalized = normalized.replace("\n ", "\n").strip()
        return normalized

    def _extract_structured_fields(self, full_text: str) -> Dict:
        """Call Gemini to identify structured legal fields from the document.
        
        Uses first 3000 + last 3000 characters to capture fields that may appear
        at the beginning or end of legal documents.
        """
        # Sample document: first 3000 + last 3000 characters
        if len(full_text) <= 6000:
            document_sample = full_text
        else:
            document_sample = full_text[:3000] + "\n\n[... document content ...]\n\n" + full_text[-3000:]
        
        prompt = (
            "Extract structured fields from this legal document.\n"
            "Return ONLY valid JSON with keys: doc_type, parties, dates, key_terms, reference_numbers.\n"
            "No explanation, no markdown fences.\n\n"
            "Document:\n"
            f"{document_sample}"
        )
        from core.genai_client import generate_with_model

        try:
            response = generate_with_model(self.model, prompt)
            raw_text = self._response_text(response)
            json_text = self._strip_json(raw_text)
            data = json.loads(json_text)
            if isinstance(data, dict):
                return {
                    "doc_type": data.get("doc_type", ""),
                    "parties": data.get("parties", []),
                    "dates": data.get("dates", []),
                    "key_terms": data.get("key_terms", []),
                    "reference_numbers": data.get("reference_numbers", []),
                }
        except Exception as e:
            print(f"⚠️  Structured field extraction error: {e}")
        
        return {
            "doc_type": "",
            "parties": [],
            "dates": [],
            "key_terms": [],
            "reference_numbers": [],
        }

    def _response_text(self, response: object) -> str:
        if isinstance(response, dict):
            return response.get("text") or response.get("content") or str(response)
        return str(response)

    def _strip_json(self, text: str) -> str:
        text = text.strip()
        text = re.sub(r"```(?:json)?", "", text, flags=re.IGNORECASE).strip()
        match = re.search(r"(\{.*\})", text, flags=re.DOTALL)
        if match:
            return match.group(1)
        return text
